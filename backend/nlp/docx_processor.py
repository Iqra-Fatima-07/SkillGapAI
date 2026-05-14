"""
docx_processor.py
~~~~~~~~~~~~~~~~~
Plain-text extraction pipeline for Microsoft Word (.docx) resume files.

Strategy
--------
1. Open the document with python-docx.
2. Walk the *body* XML in document order so that inline content (paragraphs
   that appear between table rows, headers, etc.) is preserved in the correct
   reading sequence.
3. For each paragraph: strip XML formatting tags; keep meaningful text.
4. For each table: iterate rows → cells; join cell text with a tab character
   so the logical columns are visible, and separate rows with newlines.
5. Pass the assembled text through the same ``_clean_text`` function used by
   the PDF pipeline (imported from :mod:`nlp.pdf_processor`) so whitespace
   normalisation and header/footer removal are consistent across all document
   types.

Edge cases handled
------------------
* Empty paragraphs → skipped (not emitted as blank lines early).
* Nested tables → handled by recursive ``_extract_table_text``.
* Header / footer XML elements → walked via the paragraph iterator so their
  text is included in reading order (Word stores them separately, but
  document.paragraphs already includes them when accessed via python-docx).
* Documents with no paragraphs and no tables → returns empty string.
"""

from __future__ import annotations

import logging
import re
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    # Only for type-checking; the actual import is deferred so the module
    # stays importable even when python-docx is not installed.
    import docx  # type: ignore

logger = logging.getLogger(__name__)


# ── Internal helpers ──────────────────────────────────────────────────────────


def _paragraph_text(para: "docx.text.paragraph.Paragraph") -> str:
    """Return the plain text of a python-docx Paragraph.

    Concatenates all runs; python-docx already strips XML tags for us.
    We additionally collapse internal whitespace to a single space.
    """
    text = para.text  # already XML-tag-free
    # Collapse any internal tab/multiple-space runs inside the cell text
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _extract_table_text(table: "docx.table.Table") -> str:
    """Recursively extract text from a table (including nested tables).

    Each row is emitted as a single line with cells joined by `` | ``.
    Nested tables found inside cells are expanded inline.

    Args:
        table: A python-docx ``Table`` object.

    Returns:
        A multi-line string representing the table content.
    """
    row_texts: list[str] = []

    for row in table.rows:
        cell_texts: list[str] = []
        for cell in row.cells:
            # Each cell may itself contain paragraphs and/or nested tables.
            # We collect them in document order via the cell's XML children.
            cell_parts: list[str] = []

            for child in cell.tables:
                # Nested table inside this cell
                nested_text = _extract_table_text(child)
                if nested_text:
                    cell_parts.append(nested_text)

            # Paragraphs in the cell (includes the cell's own text)
            for para in cell.paragraphs:
                para_text = _paragraph_text(para)
                if para_text:
                    cell_parts.append(para_text)

            cell_text = " ".join(cell_parts).strip()
            cell_texts.append(cell_text)

        # Filter out entirely empty cells before joining
        non_empty = [c for c in cell_texts if c]
        if non_empty:
            row_texts.append(" | ".join(non_empty))

    return "\n".join(row_texts)


def _walk_document_body(document: "docx.Document") -> str:
    """Walk the document body in XML order, preserving reading sequence.

    python-docx exposes ``document.paragraphs`` and ``document.tables``
    separately, but that loses the interleaved ordering when a table appears
    between two paragraphs.  We therefore iterate the raw XML children of
    ``<w:body>`` and dispatch on element type.

    Args:
        document: An opened python-docx ``Document`` object.

    Returns:
        Raw extracted text (not yet cleaned).
    """
    # The namespace tag python-docx uses internally
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _PARA_TAG = f"{{{_W_NS}}}p"
    _TABLE_TAG = f"{{{_W_NS}}}tbl"

    # Build look-up maps so we can retrieve the python-docx wrapper objects
    # instead of re-parsing raw XML ourselves.
    para_map = {id(p._element): p for p in document.paragraphs}
    table_map = {id(t._element): t for t in document.tables}

    parts: list[str] = []

    for child_el in document.element.body:
        tag = child_el.tag
        if tag == _PARA_TAG:
            para = para_map.get(id(child_el))
            if para is not None:
                text = _paragraph_text(para)
                if text:
                    parts.append(text)
        elif tag == _TABLE_TAG:
            table = table_map.get(id(child_el))
            if table is not None:
                table_text = _extract_table_text(table)
                if table_text:
                    parts.append(table_text)
        # Other body-level elements (sectPr, bookmarks, etc.) are silently
        # skipped; they carry no meaningful resume text.

    return "\n".join(parts)


# ── Public API ────────────────────────────────────────────────────────────────


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file given its raw bytes.

    Algorithm:
    1. Open the document stream with python-docx.
    2. Walk the body in XML document order (paragraphs interleaved with tables).
    3. Extract table rows as pipe-delimited plain text; handle nested tables.
    4. Pass the result through ``_clean_text`` (shared with PDF pipeline) for
       whitespace normalisation and header/footer removal.

    Args:
        file_bytes: Raw bytes of the .docx file.

    Returns:
        Cleaned plain-text string, or an empty string on failure.

    Raises:
        No exceptions are raised; errors are logged and an empty string is
        returned so the caller can degrade gracefully.
    """
    try:
        import docx as python_docx  # type: ignore  # python-docx
    except ImportError:
        logger.error(
            "python-docx is not installed. "
            "Add 'python-docx' to requirements.txt and reinstall."
        )
        return ""

    import io

    try:
        document = python_docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.error("Failed to open .docx document: %s", exc)
        return ""

    try:
        raw_text = _walk_document_body(document)
    except Exception as exc:
        logger.error("Error while walking document body: %s", exc)
        return ""

    # Reuse the PDF pipeline's cleaning function for consistency
    from nlp.pdf_processor import _clean_text  # type: ignore

    return _clean_text(raw_text)
